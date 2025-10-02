import functools
from flask import (
    Blueprint, flash, g, redirect, render_template, request, session, url_for, jsonify
)
from werkzeug.exceptions import abort
from bson.objectid import ObjectId
import datetime
import os
import tempfile
import google.generativeai as genai
from pptx import Presentation
import PyPDF2
import docx
from PIL import Image
import pytesseract

from flaskr.db import get_db
from flaskr.auth import login_required, recruiter_required, student_required
from flaskr.jobs import get_job
from flaskr.notifications import notify_student_shortlisted, notify_student_selected, notify_student_interview_scheduled, notify_student_interview_result
from flaskr.profile import RESUME_FOLDER

bp = Blueprint('applications', __name__, url_prefix='/applications')

@bp.route('/job/<job_id>')
@recruiter_required
def job_applications(job_id):
    """View all applications for a specific job."""
    job = get_job(job_id)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    db = get_db()
    applications = list(db['applications'].find({'job_id': ObjectId(job_id)}).sort('created_at', -1))
    
    # Add file type information for each application's resume
    for app in applications:
        if app.get('student_id'):
            student = db['students'].find_one({'_id': app['student_id']})
            if student and student.get('resume_url'):
                # Determine file type based on extension
                file_extension = student['resume_url'].rsplit('.', 1)[1].lower() if '.' in student['resume_url'] else ''
                app['resume_file_type'] = file_extension
            else:
                app['resume_file_type'] = None
        else:
            app['resume_file_type'] = None
    
    return render_template('applications/job_applications.html', job=job, applications=applications)

@bp.route('/view/<application_id>')
@recruiter_required
def view_application(application_id):
    """View detailed application with integrated PDF viewer."""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        abort(404)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    # Get the student to determine resume file type
    file_type = None
    if application.get('student_id'):
        student = db['students'].find_one({'_id': application['student_id']})
        if student and student.get('resume_url'):
            # Determine file type based on extension
            file_extension = student['resume_url'].rsplit('.', 1)[1].lower() if '.' in student['resume_url'] else ''
            file_type = file_extension
    
    return render_template('applications/application_view.html', application=application, job=job, file_type=file_type)

@bp.route('/view-pdf/<application_id>')
@recruiter_required
def view_pdf(application_id):
    """View dedicated viewer for an application's resume (supports multiple file formats)."""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        abort(404)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    # Get the student to determine resume file type
    student = db['students'].find_one({'_id': application['student_id']})
    if student is None or not student.get('resume_url'):
        flash('Resume not found', 'error')
        return redirect(url_for('applications.view_application', application_id=application_id))
    
    # Determine file type based on extension
    file_extension = student['resume_url'].rsplit('.', 1)[1].lower() if '.' in student['resume_url'] else ''
    
    return render_template('applications/pdf_viewer.html', 
                           application_id=application_id,
                           student_id=application['student_id'],
                           student_name=application['student_name'],
                           file_type=file_extension)

@bp.route('/<application_id>/update-status', methods=('POST',))
@recruiter_required
def update_status(application_id):
    """Update the status of an application."""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        abort(404)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    # Get the new status from the form
    new_status = request.form.get('status')
    if not new_status:
        flash('Status is required.', 'error')
        return redirect(url_for('applications.job_applications', job_id=str(job['_id'])))
    
    # Update the application status
    db['applications'].update_one(
        {'_id': ObjectId(application_id)},
        {'$set': {
            'status': new_status,
            'status_updated_at': datetime.datetime.now(),
            'status_updated_by': g.user['_id']
        }}
    )
    
    # Add a notification for the student
    db['notifications'].insert_one({
        'user_id': application['student_id'],
        'title': f'Application Status Updated',
        'message': f'Your application for {job["title"]} at {job["company_name"]} has been updated to: {new_status}',
        'read': False,
        'created_at': datetime.datetime.now()
    })
    
    # Get the student for SMS notification
    student = db['students'].find_one({'_id': application['student_id']})
    
    # Send SMS notification based on application status
    sms_sent = False
    if new_status == 'Shortlisted' and student:
        # Send shortlisted notification
        sms_sent = notify_student_shortlisted(student, job)
        if sms_sent:
            flash('Application status updated and SMS notification sent!', 'success')
        else:
            flash('Application status updated, but SMS notification could not be sent.', 'warning')
    elif new_status == 'Selected' and student:
        # Send selected notification
        sms_sent = notify_student_selected(student, job)
        if sms_sent:
            flash('Application status updated and SMS notification sent!', 'success')
        else:
            flash('Application status updated, but SMS notification could not be sent.', 'warning')
    else:
        flash('Application status updated successfully!', 'success')
    
    return redirect(url_for('applications.job_applications', job_id=str(job['_id'])))

@bp.route('/<application_id>/schedule-interview', methods=('GET', 'POST'))
@recruiter_required
def schedule_interview(application_id):
    """Schedule an interview for an application."""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        abort(404)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    if request.method == 'POST':
        interview_date = request.form.get('interview_date')
        interview_time = request.form.get('interview_time')
        interview_location = request.form.get('interview_location')
        interview_type = request.form.get('interview_type')
        interview_details = request.form.get('interview_details')
        
        error = None
        
        if not interview_date:
            error = 'Interview date is required.'
        elif not interview_time:
            error = 'Interview time is required.'
        elif not interview_location:
            error = 'Interview location is required.'
        elif not interview_type:
            error = 'Interview type is required.'
        
        if error is None:
            # Create a datetime object from the date and time
            interview_datetime = datetime.datetime.strptime(f'{interview_date} {interview_time}', '%Y-%m-%d %H:%M')
            
            # Create the interview
            interview_id = db['interviews'].insert_one({
                'application_id': ObjectId(application_id),
                'job_id': job['_id'],
                'student_id': application['student_id'],
                'recruiter_id': g.user['_id'],
                'interview_datetime': interview_datetime,
                'interview_location': interview_location,
                'interview_type': interview_type,
                'interview_details': interview_details,
                'status': 'Scheduled',
                'created_at': datetime.datetime.now()
            }).inserted_id
            
            # Update the application status
            db['applications'].update_one(
                {'_id': ObjectId(application_id)},
                {'$set': {
                    'status': 'Interview Scheduled',
                    'interview_id': interview_id,
                    'status_updated_at': datetime.datetime.now(),
                    'status_updated_by': g.user['_id']
                }}
            )
            
            # Add a notification for the student
            db['notifications'].insert_one({
                'user_id': application['student_id'],
                'title': f'Interview Scheduled',
                'message': f'An interview has been scheduled for your application to {job["title"]} at {job["company_name"]}. Date: {interview_date}, Time: {interview_time}',
                'read': False,
                'created_at': datetime.datetime.now()
            })
            
            # Send SMS notification if student has a phone number
            student = db['students'].find_one({'_id': application['student_id']})
            if student:
                notify_student_interview_scheduled(student, job, {
                    'interview_datetime': interview_datetime,
                    'interview_type': interview_type,
                    'interview_location': interview_location
                })
            
            flash('Interview scheduled successfully!', 'success')
            return redirect(url_for('applications.job_applications', job_id=str(job['_id'])))
        
        flash(error, 'error')
    
    # Get interview types for the form
    interview_types = [
        'In-person',
        'Phone',
        'Video',
        'Technical',
        'HR',
        'Group Discussion'
    ]
    
    return render_template('applications/schedule_interview.html', 
                          application=application, 
                          job=job, 
                          interview_types=interview_types)

@bp.route('/<application_id>/create-interview', methods=('GET', 'POST'))
@recruiter_required
def create_interview(application_id):
    """Create an interview for a selected student."""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        abort(404)
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        abort(403)
    
    # Check if the application status is 'Selected'
    if application['status'] != 'Selected':
        flash('Only selected candidates can have interviews created.', 'error')
        return redirect(url_for('applications.view_application', application_id=application_id))
    
    if request.method == 'POST':
        interview_date = request.form.get('interview_date')
        interview_time = request.form.get('interview_time')
        interview_location = request.form.get('interview_location')
        interview_type = request.form.get('interview_type')
        interview_details = request.form.get('interview_details')
        
        error = None
        
        if not interview_date:
            error = 'Interview date is required.'
        elif not interview_time:
            error = 'Interview time is required.'
        elif not interview_location:
            error = 'Interview location is required.'
        elif not interview_type:
            error = 'Interview type is required.'
        
        if error is None:
            # Create a datetime object from the date and time
            interview_datetime = datetime.datetime.strptime(f'{interview_date} {interview_time}', '%Y-%m-%d %H:%M')
            
            # Create the interview
            interview_id = db['interviews'].insert_one({
                'application_id': ObjectId(application_id),
                'job_id': job['_id'],
                'student_id': application['student_id'],
                'recruiter_id': g.user['_id'],
                'interview_datetime': interview_datetime,
                'interview_location': interview_location,
                'interview_type': interview_type,
                'interview_details': interview_details,
                'status': 'Scheduled',
                'created_at': datetime.datetime.now()
            }).inserted_id
            
            # Add a notification for the student
            db['notifications'].insert_one({
                'user_id': application['student_id'],
                'title': f'Interview Created',
                'message': f'An interview has been created for your application to {job["title"]} at {job["company_name"]}. Date: {interview_date}, Time: {interview_time}',
                'read': False,
                'created_at': datetime.datetime.now()
            })
            
            # Send SMS notification if student has a phone number
            student = db['students'].find_one({'_id': application['student_id']})
            if student:
                notify_student_interview_scheduled(student, job, {
                    'interview_datetime': interview_datetime,
                    'interview_type': interview_type,
                    'interview_location': interview_location
                })
            
            flash('Interview created successfully!', 'success')
            return redirect(url_for('applications.view_application', application_id=application_id))
        
        flash(error, 'error')
    
    # Get interview types for the form
    interview_types = [
        'In-person',
        'Phone',
        'Video',
        'Technical',
        'HR',
        'Group Discussion'
    ]
    
    return render_template('applications/create_interview.html', 
                          application=application, 
                          job=job, 
                          interview_types=interview_types)

@bp.route('/interviews')
@login_required
def interviews():
    """View all interviews for the current user."""
    db = get_db()
    
    if g.user['user_type'] == 'student':
        # Get all interviews for the student
        interviews = list(db['interviews'].find({'student_id': g.user['_id']}).sort('interview_datetime', 1))
    else:
        # Get all interviews created by the recruiter
        interviews = list(db['interviews'].find({'recruiter_id': g.user['_id']}).sort('interview_datetime', 1))
    
    # Get job and application details for each interview
    for interview in interviews:
        job = db['jobs'].find_one({'_id': interview['job_id']})
        if job:
            interview['job'] = job
        
        application = db['applications'].find_one({'_id': interview['application_id']})
        if application:
            interview['application'] = application
            
            # If recruiter, get student details
            if g.user['user_type'] == 'recruiter':
                student = db['students'].find_one({'_id': application['student_id']})
                if student:
                    interview['student'] = student
    
    # If recruiter, get all selected applications for the create interview modal
    selected_applications = []
    if g.user['user_type'] == 'recruiter':
        # Get all jobs created by this recruiter
        recruiter_jobs = list(db['jobs'].find({'recruiter_id': g.user['_id']}))
        job_ids = [job['_id'] for job in recruiter_jobs]
        
        # Get all selected applications for these jobs
        if job_ids:
            applications = list(db['applications'].find({
                'job_id': {'$in': job_ids},
                'status': 'Selected'
            }))
            
            # Add job details to each application
            for application in applications:
                job = next((j for j in recruiter_jobs if j['_id'] == application['job_id']), None)
                if job:
                    application['job'] = job
                    selected_applications.append(application)
    
    return render_template('applications/interviews.html', 
                           interviews=interviews, 
                           selected_applications=selected_applications)

@bp.route('/create-interview-from-list', methods=('POST',))
@recruiter_required
def create_interview_from_list():
    """Create an interview from the interviews list page."""
    db = get_db()
    
    # Get form data
    application_id = request.form.get('application_id')
    interview_date = request.form.get('interview_date')
    interview_time = request.form.get('interview_time')
    interview_location = request.form.get('interview_location')
    interview_type = request.form.get('interview_type')
    interview_details = request.form.get('interview_details')
    
    error = None
    
    if not application_id:
        error = 'Student is required.'
    elif not interview_date:
        error = 'Interview date is required.'
    elif not interview_time:
        error = 'Interview time is required.'
    elif not interview_location:
        error = 'Interview location is required.'
    elif not interview_type:
        error = 'Interview type is required.'
    
    if error is None:
        # Get the application
        application = db['applications'].find_one({'_id': ObjectId(application_id)})
        if application is None:
            abort(404)
        
        # Get the job
        job = db['jobs'].find_one({'_id': application['job_id']})
        if job is None:
            abort(404)
        
        # Check if the current user is the creator of this job listing
        if g.user['_id'] != job['recruiter_id']:
            abort(403)
        
        # Create a datetime object from the date and time
        interview_datetime = datetime.datetime.strptime(f'{interview_date} {interview_time}', '%Y-%m-%d %H:%M')
        
        # Create the interview
        interview_id = db['interviews'].insert_one({
            'application_id': ObjectId(application_id),
            'job_id': job['_id'],
            'student_id': application['student_id'],
            'recruiter_id': g.user['_id'],
            'interview_datetime': interview_datetime,
            'interview_location': interview_location,
            'interview_type': interview_type,
            'interview_details': interview_details,
            'status': 'Scheduled',
            'created_at': datetime.datetime.now()
        }).inserted_id
        
        # Add a notification for the student
        db['notifications'].insert_one({
            'user_id': application['student_id'],
            'title': f'Interview Created',
            'message': f'An interview has been created for your application to {job["title"]} at {job["company_name"]}. Date: {interview_date}, Time: {interview_time}',
            'read': False,
            'created_at': datetime.datetime.now()
        })
        
        # Send SMS notification if student has a phone number
        student = db['students'].find_one({'_id': application['student_id']})
        if student:
            notify_student_interview_scheduled(student, job, {
                'interview_datetime': interview_datetime,
                'interview_type': interview_type,
                'interview_location': interview_location
            })
        
        flash('Interview created successfully!', 'success')
    else:
        flash(error, 'error')
    
    return redirect(url_for('applications.interviews'))

@bp.route('/interview/<interview_id>/view')
@login_required
def interview_view(interview_id):
    """View details of an interview."""
    db = get_db()
    
    # Get the interview
    interview = db['interviews'].find_one({'_id': ObjectId(interview_id)})
    if interview is None:
        abort(404)
    
    # Check if the current user is authorized to view this interview
    if g.user['user_type'] == 'student' and g.user['_id'] != interview['student_id']:
        abort(403)
    elif g.user['user_type'] == 'recruiter' and g.user['_id'] != interview['recruiter_id']:
        abort(403)
    
    # Get job and application details
    job = db['jobs'].find_one({'_id': interview['job_id']})
    if job:
        interview['job'] = job
    
    application = db['applications'].find_one({'_id': interview['application_id']})
    if application:
        interview['application'] = application
    
    return render_template('applications/interview_view.html', interview=interview)

@bp.route('/interview/<interview_id>/result')
@recruiter_required
def interview_result(interview_id):
    """Show the form to update an interview result."""
    db = get_db()
    
    # Get the interview
    interview = db['interviews'].find_one({'_id': ObjectId(interview_id)})
    if interview is None:
        abort(404)
    
    # Check if the current user is the creator of this interview
    if g.user['_id'] != interview['recruiter_id']:
        abort(403)
    
    # Check if the interview is still scheduled
    if interview['status'] != 'Scheduled':
        flash('This interview has already been completed or cancelled.', 'error')
        return redirect(url_for('applications.interview_view', interview_id=interview_id))
    
    # Get job and application details
    job = db['jobs'].find_one({'_id': interview['job_id']})
    if job:
        interview['job'] = job
    
    application = db['applications'].find_one({'_id': interview['application_id']})
    if application:
        interview['application'] = application
    
    return render_template('applications/interview_result.html', interview=interview)

@bp.route('/interview/<interview_id>/update-result', methods=('POST',))
@recruiter_required
def update_interview_result(interview_id):
    """Update the result of an interview."""
    db = get_db()
    
    # Get the interview
    interview = db['interviews'].find_one({'_id': ObjectId(interview_id)})
    if interview is None:
        abort(404)
    
    # Check if the current user is the creator of this interview
    if g.user['_id'] != interview['recruiter_id']:
        abort(403)
    
    # Get the application
    application = db['applications'].find_one({'_id': interview['application_id']})
    if application is None:
        abort(404)
    
    # Get the job
    job = db['jobs'].find_one({'_id': interview['job_id']})
    if job is None:
        abort(404)
    
    # Get the result from the form
    result = request.form.get('result')
    feedback = request.form.get('feedback', '')
    
    if not result:
        flash('Result is required.', 'error')
        return redirect(url_for('applications.interview_result', interview_id=interview_id))
    
    # Update the interview status
    db['interviews'].update_one(
        {'_id': ObjectId(interview_id)},
        {'$set': {
            'status': 'Completed',
            'result': result,
            'feedback': feedback,
            'completed_at': datetime.datetime.now(),
            'completed_by': g.user['_id']
        }}
    )
    
    # Update the application status based on the result
    new_status = 'Selected' if result == 'Pass' else 'Rejected'
    
    db['applications'].update_one(
        {'_id': interview['application_id']},
        {'$set': {
            'status': new_status,
            'status_updated_at': datetime.datetime.now(),
            'status_updated_by': g.user['_id']
        }}
    )
    
    # Add a notification for the student
    db['notifications'].insert_one({
        'user_id': application['student_id'],
        'title': f'Interview Result: {result}',
        'message': f'Your interview for {job["title"]} at {job["company_name"]} has been marked as {result}. {feedback}',
        'read': False,
        'created_at': datetime.datetime.now()
    })
    
    # Send SMS notification if student has a phone number
    student = db['students'].find_one({'_id': application['student_id']})
    if student:
        notify_student_interview_result(student, job, {
            'result': result
        })
        
        # If the student is selected, also send the selection notification
        if result == 'Pass':
            notify_student_selected(student, job)
    
    flash('Interview result updated successfully!', 'success')
    return redirect(url_for('applications.interview_view', interview_id=interview_id))

@bp.route('/notifications')
@login_required
def notifications():
    """View all notifications for the current user."""
    db = get_db()
    
    # Get all notifications for the user
    notifications = list(db['notifications'].find({'user_id': g.user['_id']}).sort('created_at', -1))
    
    # Mark all unread notifications as read
    db['notifications'].update_many(
        {'user_id': g.user['_id'], 'read': False},
        {'$set': {'read': True}}
    )
    
    return render_template('applications/notifications.html', notifications=notifications)


# Configure Google Gemini API
# For production, use environment variables instead
GEMINI_API_KEY = "your gemi api key here"  # Replace with your actual Gemini API key
genai.configure(api_key=GEMINI_API_KEY)


def extract_text_from_pdf(file_path):
    """Extract text content from a PDF file"""
    text_content = []
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
        
        return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""


def extract_text_from_docx(file_path):
    """Extract text content from a Word document"""
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return ""


def extract_text_from_image(file_path):
    """Extract text content from an image using OCR"""
    try:
        # Open the image
        image = Image.open(file_path)
        
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(image)
        
        return text
    except Exception as e:
        print(f"Error extracting text from image: {str(e)}")
        return ""


def generate_resume_summary(text, job_title=None, job_description=None):
    """Generate a summary of the resume using Google Gemini API"""
    if not text.strip():
        return {
            "candidate_summary": "<p>No text content could be extracted from the resume.</p>",
            "key_skills": "<p>No skills could be identified.</p>",
            "job_fit": "<p>Unable to analyze job fit due to missing resume content.</p>"
        }

    try:
        # Use Gemini model
        model = genai.GenerativeModel("models/gemini-1.5-flash")
        
        # Create a prompt for resume analysis
        job_context = ""
        if job_title and job_description:
            job_context = f"\nThe candidate has applied for the position of {job_title}. \nJob Description: {job_description}\n"
        
        prompt = f"""
        You are an expert HR professional analyzing a resume. Please provide a comprehensive analysis of the following resume content.
        Format your response in HTML with appropriate tags (<p>, <ul>, <li>, <strong>, etc.).
        
        Resume Content:
        {text}
        {job_context}
        
        Please provide the following sections:
        1. Candidate Summary: A brief overview of the candidate's background, experience, and qualifications.
        2. Key Skills: A bullet-point list of the candidate's key skills and competencies.
        3. Job Fit Analysis: An assessment of how well the candidate's profile matches the job requirements.
        
        Return your response as a JSON object with the following structure:
        {{
            "candidate_summary": "HTML formatted candidate summary",
            "key_skills": "HTML formatted list of key skills",
            "job_fit": "HTML formatted job fit analysis"
        }}
        """
        
        # Generate the response
        response = model.generate_content(prompt)
        
        # Try to parse the response as JSON
        try:
            import json
            # Extract JSON from the response text
            response_text = response.text
            
            # Find JSON content between curly braces
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            
            if start_idx >= 0 and end_idx > start_idx:
                json_str = response_text[start_idx:end_idx]
                result = json.loads(json_str)
                return result
            else:
                # Fallback: create structured response manually
                return {
                    "candidate_summary": f"<p>{response.text}</p>",
                    "key_skills": "<p>Skills extraction failed.</p>",
                    "job_fit": "<p>Job fit analysis failed.</p>"
                }
        except Exception as json_error:
            print(f"Error parsing JSON response: {str(json_error)}")
            # Fallback: return the raw text
            return {
                "candidate_summary": f"<p>{response.text}</p>",
                "key_skills": "<p>Skills extraction failed.</p>",
                "job_fit": "<p>Job fit analysis failed.</p>"
            }
    
    except Exception as e:
        print(f"Error generating summary: {str(e)}")
        # Try fallback to another model if the first one fails
        try:
            fallback_model = genai.GenerativeModel("models/gemini-1.5-pro")
            response = fallback_model.generate_content(prompt)
            return {
                "candidate_summary": f"<p>{response.text}</p>",
                "key_skills": "<p>Skills extraction failed.</p>",
                "job_fit": "<p>Job fit analysis failed.</p>"
            }
        except Exception as fallback_e:
            error_message = f"Error generating summary with primary model: {str(e)}\n\nError with fallback model: {str(fallback_e)}"
            return {
                "candidate_summary": f"<p>Error generating summary: {error_message}</p>",
                "key_skills": "<p>Skills extraction failed.</p>",
                "job_fit": "<p>Job fit analysis failed.</p>"
            }


@bp.route('/resume-summary/<application_id>')
@recruiter_required
def resume_summary(application_id):
    """Generate an AI summary of a student's resume"""
    db = get_db()
    
    # Get the application
    application = db['applications'].find_one({'_id': ObjectId(application_id)})
    if application is None:
        return jsonify({
            'error': 'Application not found'
        }), 404
    
    # Get the job
    job = db['jobs'].find_one({'_id': application['job_id']})
    if job is None:
        return jsonify({
            'error': 'Job not found'
        }), 404
    
    # Check if the current user is the creator of this job listing
    if g.user['_id'] != job['recruiter_id']:
        return jsonify({
            'error': 'Unauthorized access'
        }), 403
    
    # Get the student to determine resume file type and path
    student = db['students'].find_one({'_id': application['student_id']})
    if student is None or not student.get('resume_url'):
        return jsonify({
            'error': 'Resume not found'
        }), 404
    
    # Get the resume file path
    resume_path = os.path.join(RESUME_FOLDER, student['resume_url'])
    
    if not os.path.exists(resume_path):
        return jsonify({
            'error': 'Resume file not found'
        }), 404
    
    # Determine file type based on extension
    file_extension = student['resume_url'].rsplit('.', 1)[1].lower() if '.' in student['resume_url'] else ''
    
    # Extract text based on file type
    text_content = ""
    if file_extension == 'pdf':
        text_content = extract_text_from_pdf(resume_path)
    elif file_extension in ['doc', 'docx']:
        text_content = extract_text_from_docx(resume_path)
    elif file_extension in ['jpg', 'jpeg']:
        text_content = extract_text_from_image(resume_path)
    else:
        return jsonify({
            'error': 'Unsupported file type'
        }), 400
    
    # Generate summary using the extracted text
    summary = generate_resume_summary(
        text_content, 
        job_title=job.get('title'), 
        job_description=job.get('description')
    )
    
    return jsonify(summary)
